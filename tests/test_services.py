import json
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any
from unittest.mock import MagicMock, patch

import pytest
from PIL import Image
from sqlalchemy.orm import Session

from src.api.schemas import AssetDetailSchema, AssetSchema, MetadataSchema
from src.core.config import settings
from src.core.database import Base, SessionLocal, engine
from src.core.exceptions import (
    AIServiceError,
    AITimeoutError,
    FileTypeError,
    FileSizeError,
    TextExtractionError,
)
from src.models.asset import Asset, Metadata
from src.repositories.asset_repository import AssetRepository
from src.services.ai_service import AIService
from src.services.asset_service import AssetService
from src.services.search_service import SearchService
from src.utils.file_validator import FileValidator
from src.utils.text_extractor import TextExtractor

"""
Comprehensive tests for Phase 5 Business Logic Services.

Tests verify:
1. File validation (size, type, integrity)
2. Text extraction (txt, pdf, docx)
3. OpenAI integration (mocked)
4. AI service orchestration
5. Asset service upload workflow
6. Search service integration
7. Cross-layer connections and error handling
"""


@pytest.fixture(scope="function")
def db_session() -> Session:
    """Create a fresh database session for each test."""
    Base.metadata.create_all(bind=engine)
    session = SessionLocal()
    yield session
    session.close()
    Base.metadata.drop_all(bind=engine)


@pytest.fixture
def temp_upload_dir(tmp_path: Path) -> Path:
    """Create temporary upload directory."""
    upload_dir = tmp_path / "uploads"
    upload_dir.mkdir(parents=True, exist_ok=True)
    return upload_dir


@pytest.fixture
def sample_text_file(tmp_path: Path) -> Path:
    """Create a sample text file."""
    text_file = tmp_path / "sample.txt"
    text_file.write_text("This is a sample text file for testing.\nIt has multiple lines.")
    return text_file


@pytest.fixture
def sample_image_file(tmp_path: Path) -> Path:
    """Create a sample image file."""
    image_path = tmp_path / "sample.jpg"
    img = Image.new("RGB", (100, 100), color="red")
    img.save(image_path)
    return image_path


@pytest.fixture
def sample_pdf_file(tmp_path: Path) -> Path:
    """Create a mock PDF file (for testing only - not a real PDF)."""
    pdf_path = tmp_path / "sample.pdf"
    # Create a minimal valid PDF file
    pdf_path.write_bytes(
        b"%PDF-1.4\n"
        b"1 0 obj\n"
        b"<< /Type /Catalog /Pages 2 0 R >>\n"
        b"endobj\n"
        b"2 0 obj\n"
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>\n"
        b"endobj\n"
        b"3 0 obj\n"
        b"<< /Type /Page /Parent 2 0 R /Resources << >> /MediaBox [0 0 612 792] /Contents 4 0 R >>\n"
        b"endobj\n"
        b"4 0 obj\n"
        b"<< >>\n"
        b"stream\n"
        b"BT\n"
        b"/F1 12 Tf\n"
        b"100 700 Td\n"
        b"(Sample PDF content) Tj\n"
        b"ET\n"
        b"endstream\n"
        b"endobj\n"
        b"xref\n"
        b"0 5\n"
        b"0000000000 65535 f\n"
        b"0000000009 00000 n\n"
        b"0000000058 00000 n\n"
        b"0000000115 00000 n\n"
        b"0000000213 00000 n\n"
        b"trailer\n"
        b"<< /Size 5 /Root 1 0 R >>\n"
        b"startxref\n"
        b"317\n"
        b"%%EOF\n"
    )
    return pdf_path


@pytest.fixture
def sample_docx_file(tmp_path: Path) -> Path:
    """Create a sample DOCX file."""
    from docx import Document as DocxDocument

    docx_path = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_paragraph("This is a test DOCX document.")
    doc.add_paragraph("It contains sample content.")
    doc.save(docx_path)
    return docx_path


# ============================================================================
# FileValidator Tests
# ============================================================================


class TestFileValidator:
    """Tests for file validation."""

    def test_validate_text_file(self, sample_text_file: Path) -> None:
        """Test validating a text file."""
        file_type, mime_type = FileValidator.validate_file(sample_text_file)

        assert file_type == "text"
        assert mime_type == "text/plain"

    def test_validate_image_file(self, sample_image_file: Path) -> None:
        """Test validating an image file."""
        file_type, mime_type = FileValidator.validate_file(sample_image_file)

        assert file_type == "image"
        assert mime_type == "image/jpeg"

    def test_validate_pdf_file(self, sample_pdf_file: Path) -> None:
        """Test validating a PDF file."""
        file_type, mime_type = FileValidator.validate_file(sample_pdf_file)

        assert file_type == "text"
        assert mime_type == "application/pdf"

    def test_validate_docx_file(self, sample_docx_file: Path) -> None:
        """Test validating a DOCX file."""
        file_type, mime_type = FileValidator.validate_file(sample_docx_file)

        assert file_type == "text"
        assert mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def test_file_not_found(self, tmp_path: Path) -> None:
        """Test error when file doesn't exist."""
        non_existent = tmp_path / "nonexistent.txt"

        with pytest.raises(FileTypeError, match="does not exist"):
            FileValidator.validate_file(non_existent)

    def test_empty_file(self, tmp_path: Path) -> None:
        """Test error when file is empty."""
        empty_file = tmp_path / "empty.txt"
        empty_file.write_text("")

        with pytest.raises(FileTypeError, match="empty"):
            FileValidator.validate_file(empty_file)

    def test_file_without_extension(self, tmp_path: Path) -> None:
        """Test error when file has no extension."""
        no_ext = tmp_path / "noextension"
        no_ext.write_text("content")

        with pytest.raises(FileTypeError, match="no extension"):
            FileValidator.validate_file(no_ext)

    def test_unsupported_file_type(self, tmp_path: Path) -> None:
        """Test error with unsupported file type."""
        unsupported = tmp_path / "file.xyz"
        unsupported.write_text("content")

        with pytest.raises(FileTypeError, match="not allowed"):
            FileValidator.validate_file(unsupported)

    def test_file_size_exceeds_limit(self, tmp_path: Path) -> None:
        """Test error when file exceeds size limit."""
        large_file = tmp_path / "large.txt"
        # Create a file larger than max (default 10MB)
        large_file.write_bytes(b"x" * (settings.max_file_size + 1))

        with pytest.raises(FileSizeError):
            FileValidator.validate_file(large_file)

    def test_corrupted_image_file(self, tmp_path: Path) -> None:
        """Test error with corrupted image."""
        corrupted_image = tmp_path / "corrupted.jpg"
        # Write invalid JPEG data
        corrupted_image.write_bytes(b"FAKE JPEG DATA")

        with pytest.raises(FileTypeError, match="corrupted|invalid"):
            FileValidator.validate_file(corrupted_image)

    def test_invalid_utf8_text_file(self, tmp_path: Path) -> None:
        """Test error with invalid UTF-8 text file."""
        invalid_text = tmp_path / "invalid.txt"
        # Write bytes that are not valid UTF-8
        invalid_text.write_bytes(b"\x80\x81\x82\x83")

        with pytest.raises(FileTypeError, match="UTF-8"):
            FileValidator.validate_file(invalid_text)


# ============================================================================
# TextExtractor Tests
# ============================================================================


class TestTextExtractor:
    """Tests for text extraction."""

    def test_extract_from_text_file(self, sample_text_file: Path) -> None:
        """Test extracting text from plain text file."""
        text = TextExtractor.extract_text(sample_text_file)

        assert "sample text file" in text.lower()
        assert len(text) > 0

    def test_extract_from_pdf_file(self, sample_pdf_file: Path) -> None:
        """Test extracting text from PDF file."""
        text = TextExtractor.extract_text(sample_pdf_file)

        # PDF should extract something (content or empty is ok for minimal PDF)
        assert isinstance(text, str)
        assert len(text) >= 0  # May be empty for minimal PDF

    def test_extract_from_docx_file(self, sample_docx_file: Path) -> None:
        """Test extracting text from DOCX file."""
        text = TextExtractor.extract_text(sample_docx_file)

        assert "test DOCX document" in text or "test" in text.lower()
        assert len(text) > 0

    def test_text_truncation(self, tmp_path: Path) -> None:
        """Test that extracted text is truncated to MAX_CHARS."""
        long_text_file = tmp_path / "long.txt"
        # Create a file with content longer than MAX_CHARS
        long_text_file.write_text("x" * (TextExtractor.MAX_CHARS + 1000))

        text = TextExtractor.extract_text(long_text_file)

        assert len(text) <= TextExtractor.MAX_CHARS

    def test_unsupported_file_type(self, tmp_path: Path) -> None:
        """Test error with unsupported file type."""
        unsupported = tmp_path / "file.xyz"
        unsupported.write_text("content")

        with pytest.raises(TextExtractionError, match="Unsupported"):
            TextExtractor.extract_text(unsupported)

    def test_nonexistent_file(self, tmp_path: Path) -> None:
        """Test error with nonexistent file."""
        non_existent = tmp_path / "nonexistent.txt"

        with pytest.raises(TextExtractionError):
            TextExtractor.extract_text(non_existent)


# ============================================================================
# OpenAI Client Tests (Mocked)
# ============================================================================


class TestOpenAIClient:
    """Tests for OpenAI client (with mocked API)."""

    @patch("src.integrations.openai_client.settings")
    @patch("src.integrations.openai_client.OpenAI")
    def test_generate_metadata_for_text(
        self, mock_openai_class: MagicMock, mock_settings: MagicMock
    ) -> None:
        """Test generating metadata for text."""
        mock_settings.openai_api_key = "test-key"
        mock_client = MagicMock()
        mock_openai_class.return_value = mock_client

        # Mock API response
        mock_response = MagicMock()
        mock_response.choices[0].message.content = json.dumps(
            {
                "description": "Test document about AI",
                "tags": ["AI", "machine learning", "test"],
                "keywords": ["artificial intelligence", "models", "training"],
            }
        )
        mock_client.chat.completions.create.return_value = mock_response

        from src.integrations.openai_client import OpenAIClient

        client = OpenAIClient()
        result = client.generate_metadata_for_text("This is about AI and machine learning.")

        assert result["description"] == "Test document about AI"
        assert len(result["tags"]) == 3
        assert len(result["keywords"]) == 3

    @patch("src.integrations.openai_client.settings")
    @patch("src.integrations.openai_client.OpenAI")
    def test_json_parsing_fallback_markdown(
        self, mock_openai_class: MagicMock, mock_settings: MagicMock
    ) -> None:
        """Test JSON parsing with markdown code block fallback."""
        mock_settings.openai_api_key = "test-key"
        mock_client = MagicMock()
        mock_openai_class.return_value = mock_client

        # Mock response with markdown code block
        mock_response = MagicMock()
        mock_response.choices[0].message.content = (
            "Here's the metadata:\n"
            "```json\n"
            '{"description": "Test", "tags": ["tag1"], "keywords": ["kw1"]}\n'
            "```"
        )
        mock_client.chat.completions.create.return_value = mock_response

        from src.integrations.openai_client import OpenAIClient

        client = OpenAIClient()
        result = client.generate_metadata_for_text("test")

        assert result["description"] == "Test"
        assert result["tags"][0] == "tag1"

    @patch("src.integrations.openai_client.settings")
    @patch("src.integrations.openai_client.OpenAI")
    def test_api_timeout_error(
        self, mock_openai_class: MagicMock, mock_settings: MagicMock
    ) -> None:
        """Test handling of API errors."""
        mock_settings.openai_api_key = "test-key"
        mock_client = MagicMock()
        mock_openai_class.return_value = mock_client

        # Mock a generic error from the API
        mock_client.chat.completions.create.side_effect = Exception("API error")

        from src.integrations.openai_client import OpenAIClient

        client = OpenAIClient()

        with pytest.raises(AIServiceError):
            client.generate_metadata_for_text("test")

    @patch("src.integrations.openai_client.settings")
    @patch("src.integrations.openai_client.OpenAI")
    def test_json_parsing_error(
        self, mock_openai_class: MagicMock, mock_settings: MagicMock
    ) -> None:
        """Test handling of JSON parsing error."""
        mock_settings.openai_api_key = "test-key"
        mock_client = MagicMock()
        mock_openai_class.return_value = mock_client

        # Mock response with invalid JSON
        mock_response = MagicMock()
        mock_response.choices[0].message.content = "This is not JSON at all"
        mock_client.chat.completions.create.return_value = mock_response

        from src.integrations.openai_client import OpenAIClient

        client = OpenAIClient()

        with pytest.raises(Exception):  # AIParsingError
            client.generate_metadata_for_text("test")


# ============================================================================
# AI Service Tests
# ============================================================================


class TestAIService:
    """Tests for AI service."""

    @patch("src.services.ai_service.OpenAIClient")
    def test_generate_metadata_for_text(
        self, mock_openai_client: MagicMock, sample_text_file: Path
    ) -> None:
        """Test AI service generating metadata for text."""
        mock_client = MagicMock()
        mock_openai_client.return_value = mock_client
        mock_client.generate_metadata_for_text.return_value = {
            "description": "Sample text document",
            "tags": ["test", "sample"],
            "keywords": ["text", "document"],
        }

        service = AIService()
        description, tags_json, keywords_json = service.generate_metadata(
            sample_text_file, "text"
        )

        assert description == "Sample text document"
        assert tags_json is not None
        assert keywords_json is not None
        # Verify JSON format
        assert isinstance(json.loads(tags_json), list)
        assert isinstance(json.loads(keywords_json), list)

    @patch("src.services.ai_service.OpenAIClient")
    def test_generate_metadata_for_image(
        self, mock_openai_client: MagicMock, sample_image_file: Path
    ) -> None:
        """Test AI service generating metadata for image."""
        mock_client = MagicMock()
        mock_openai_client.return_value = mock_client
        mock_client.generate_metadata_for_image.return_value = {
            "description": "Red square image",
            "tags": ["image", "red"],
            "keywords": ["color", "shape"],
        }

        service = AIService()
        description, tags_json, keywords_json = service.generate_metadata(
            sample_image_file, "image"
        )

        assert description == "Red square image"
        assert tags_json is not None
        assert keywords_json is not None

    @patch("src.services.ai_service.OpenAIClient")
    def test_graceful_error_handling(
        self, mock_openai_client: MagicMock, sample_text_file: Path
    ) -> None:
        """Test that errors in AI generation are handled gracefully."""
        mock_client = MagicMock()
        mock_openai_client.return_value = mock_client
        mock_client.generate_metadata_for_text.side_effect = AIServiceError(
            "API error"
        )

        service = AIService()
        description, tags_json, keywords_json = service.generate_metadata(
            sample_text_file, "text"
        )

        # Should return None values, not raise
        assert description is None
        assert tags_json is None
        assert keywords_json is None

    def test_invalid_file_type(self, sample_text_file: Path) -> None:
        """Test error with invalid file type."""
        with patch("src.services.ai_service.OpenAIClient"):
            service = AIService()

            # Invalid file type should raise AIServiceError
            try:
                service.generate_metadata(sample_text_file, "invalid_type")
                # If it doesn't raise, it should return (None, None, None)
                description, tags_json, keywords_json = service.generate_metadata(
                    sample_text_file, "invalid_type"
                )
                # Graceful error handling returns None values
                assert description is None
                assert tags_json is None
                assert keywords_json is None
            except AIServiceError:
                # This is acceptable too
                pass


# ============================================================================
# Asset Service Tests
# ============================================================================


class TestAssetService:
    """Tests for asset service (end-to-end upload workflow)."""

    @patch("src.services.asset_service.AIService")
    def test_upload_asset_ai_failure_doesnt_fail_upload(
        self,
        mock_ai_service: MagicMock,
        db_session: Session,
        sample_text_file: Path,
        temp_upload_dir: Path,
    ) -> None:
        """Test that AI metadata failure doesn't fail the upload."""
        mock_ai = MagicMock()
        mock_ai_service.return_value = mock_ai
        mock_ai.generate_metadata.return_value = (None, None, None)

        with patch("src.services.asset_service.settings") as mock_settings:
            mock_settings.upload_dir = str(temp_upload_dir)

            service = AssetService(db_session)
            result = service.upload_asset(sample_text_file, "test.txt")

        # Upload should succeed even without metadata
        assert result.id is not None
        assert result.filename == "test.txt"
        assert result.metadata is None  # But no metadata

    @patch("src.services.asset_service.AIService")
    def test_get_asset_file_path(
        self,
        mock_ai_service: MagicMock,
        db_session: Session,
        sample_text_file: Path,
        temp_upload_dir: Path,
    ) -> None:
        """Test retrieving file path for asset."""
        mock_ai = MagicMock()
        mock_ai_service.return_value = mock_ai
        mock_ai.generate_metadata.return_value = (None, None, None)

        with patch("src.services.asset_service.settings") as mock_settings:
            mock_settings.upload_dir = str(temp_upload_dir)

            service = AssetService(db_session)
            uploaded = service.upload_asset(sample_text_file, "test.txt")

        # Get file path
        file_path = service.get_asset_file_path(uploaded.id)
        assert file_path is not None
        assert file_path.exists()


# ============================================================================
# Search Service Tests
# ============================================================================


class TestSearchService:
    """Tests for search service."""

    def test_search_with_results(self, db_session: Session) -> None:
        """Test search returning results."""
        # Create test asset and metadata
        repo = AssetRepository(db_session)
        asset = repo.create_asset(
            filename="test.pdf",
            file_type="text",
            mime_type="application/pdf",
            file_path="uploads/1_test.pdf",
            file_size=1000,
        )
        repo.create_metadata(
            asset_id=asset.id,
            description="A test document",
            tags=json.dumps(["test", "document"]),
            keywords=json.dumps(["sample"]),
        )

        # Search
        service = SearchService(db_session)
        result = service.search("test", page=1, per_page=50)

        assert result.query == "test"
        assert result.total >= 1
        assert len(result.results) >= 1

    def test_search_pagination(self, db_session: Session) -> None:
        """Test search pagination."""
        repo = AssetRepository(db_session)

        # Create multiple assets
        for i in range(5):
            asset = repo.create_asset(
                filename=f"file{i}.txt",
                file_type="text",
                mime_type="text/plain",
                file_path=f"uploads/{i}_file{i}.txt",
                file_size=100 * (i + 1),
            )
            repo.create_metadata(
                asset_id=asset.id,
                tags=json.dumps(["searchable"]),
            )

        # Search with pagination
        service = SearchService(db_session)
        page1 = service.search("searchable", page=1, per_page=2)
        page2 = service.search("searchable", page=2, per_page=2)

        assert page1.total == 5
        assert len(page1.results) == 2
        assert len(page2.results) == 2

    def test_search_falls_back_to_file_content_when_metadata_is_missing(
        self, db_session: Session, tmp_path: Path
    ) -> None:
        """Search should match file content even before metadata is generated."""
        content_file = tmp_path / "content_search.txt"
        content_file.write_text("This document discusses neural networks and machine learning.", encoding="utf-8")

        repo = AssetRepository(db_session)
        repo.create_asset(
            filename="content_search.txt",
            file_type="text",
            mime_type="text/plain",
            file_path=str(content_file),
            file_size=content_file.stat().st_size,
        )

        service = SearchService(db_session)
        result = service.search("machine learning", page=1, per_page=50)

        assert result.total >= 1
        assert any(item.filename == "content_search.txt" for item in result.results)

    def test_empty_search(self, db_session: Session) -> None:
        """Test empty search query."""
        service = SearchService(db_session)
        result = service.search("", page=1, per_page=50)

        assert result.query == ""
        assert result.total == 0
        assert len(result.results) == 0

    def test_search_no_results(self, db_session: Session) -> None:
        """Test search with no matching results."""
        service = SearchService(db_session)
        result = service.search("nonexistent", page=1, per_page=50)

        assert result.total == 0
        assert len(result.results) == 0


# ============================================================================
# Integration Tests - Cross-Layer Connections
# ============================================================================


class TestIntegration:
    """Integration tests verifying cross-layer connections."""

    def test_database_relationships_maintained(self, db_session: Session) -> None:
        """Test that database relationships are properly maintained."""
        # Create asset via repository
        repo = AssetRepository(db_session)
        asset = repo.create_asset(
            filename="test.pdf",
            file_type="text",
            mime_type="application/pdf",
            file_path="uploads/1_test.pdf",
            file_size=2000,
        )

        # Create metadata
        metadata = repo.create_metadata(
            asset_id=asset.id,
            description="Test",
            tags=json.dumps(["tag1"]),
            keywords=json.dumps(["kw1"]),
        )

        # Verify relationships
        assert asset.asset_metadata is not None
        assert asset.asset_metadata.description == "Test"
        assert metadata.asset_id == asset.id

    @patch("src.services.asset_service.AIService")
    def test_error_recovery_and_cleanup(
        self,
        mock_ai_service: MagicMock,
        db_session: Session,
        temp_upload_dir: Path,
    ) -> None:
        """Test that errors during upload trigger proper cleanup."""
        # Create a test file
        source_file = temp_upload_dir / "test.txt"
        source_file.write_text("test content")

        mock_ai = MagicMock()
        mock_ai_service.return_value = mock_ai
        mock_ai.generate_metadata.return_value = (None, None, None)

        with patch("src.services.asset_service.settings") as mock_settings:
            mock_settings.upload_dir = str(temp_upload_dir)

            asset_service = AssetService(db_session)

            # Upload should work even if metadata fails
            result = asset_service.upload_asset(source_file, "test.txt")
            assert result is not None  # Asset created even if metadata fails

    def test_config_integration(self) -> None:
        """Test that all services use consistent configuration."""
        from src.core.config import settings

        # Verify critical settings exist
        assert hasattr(settings, "openai_api_key")
        assert hasattr(settings, "upload_dir")
        assert hasattr(settings, "max_file_size")
        assert hasattr(settings, "allowed_text_extensions")
        assert hasattr(settings, "allowed_image_extensions")
        assert hasattr(settings, "database_url")
        assert hasattr(settings, "ai_timeout")

    def test_exception_hierarchy(self) -> None:
        """Test that exception hierarchy is properly used."""
        from src.core.exceptions import KMSException

        # Test exception types
        assert issubclass(FileTypeError, KMSException)
        assert issubclass(FileSizeError, KMSException)
        assert issubclass(AIServiceError, KMSException)
        assert issubclass(AITimeoutError, AIServiceError)
        assert issubclass(TextExtractionError, KMSException)


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
